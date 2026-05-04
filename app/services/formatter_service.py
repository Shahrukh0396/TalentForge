import re
from pathlib import Path
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SECTION_ORDER = [
    "SUMMARY OF QUALIFICATIONS",
    "PROFESSIONAL EXPERIENCE",
    "EDUCATION",
    "TECHNICAL SKILLS",
]

SECTION_ALIASES = {
    "SUMMARY": "SUMMARY OF QUALIFICATIONS",
    "SUMMARY OF QUALIFICATIONS": "SUMMARY OF QUALIFICATIONS",
    "PROFESSIONAL EXPERIENCE": "PROFESSIONAL EXPERIENCE",
    "EXPERIENCE": "PROFESSIONAL EXPERIENCE",
    "WORK EXPERIENCE": "PROFESSIONAL EXPERIENCE",
    "EDUCATION": "EDUCATION",
    "TECHNICAL SKILLS": "TECHNICAL SKILLS",
    "SKILLS": "TECHNICAL SKILLS",
}

SOFTWARE_NORMALIZATION = {
    "quickbooks": "QuickBooks",
    "peoplesoft": "PeopleSoft",
    "powerpoint": "PowerPoint",
    "salesforce": "Salesforce",
    "microstrategy": "MicroStrategy",
}

DATE_PATTERN = re.compile(r"\b\d{2}/\d{4}\b")
PHONE_LINE_PATTERN = re.compile(
    r"(?:\+?\d{1,3}[\s\-.]?)?(?:\(\d{2,4}\)|\d{2,4})[\s\-.]?\d{3}[\s\-.]?\d{4}\b"
)
MONTH_DATE_RANGE_PATTERN = re.compile(
    r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*[,\s]+\d{4}\s*[–-]\s*(?:present|current|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*[,\s]+\d{4})\b",
    flags=re.IGNORECASE,
)


def _is_probable_name(line: str) -> bool:
    words = [word for word in re.split(r"\s+", line.strip()) if word]
    if not words or len(words) > 6:
        return False
    return all(re.fullmatch(r"[A-Za-z.\-']+", word) for word in words)


def _normalize_software_names(line: str) -> str:
    out = line
    for raw, normalized in SOFTWARE_NORMALIZATION.items():
        out = re.sub(rf"\b{raw}\b", normalized, out, flags=re.IGNORECASE)
    return out


def _strip_contact_line(line: str) -> bool:
    lowered = line.lower()
    if any(token in lowered for token in ["linkedin", "@", "www.", "http://", "https://"]):
        return True

    # Remove only phone-like patterns; do not treat date ranges as phone numbers.
    if PHONE_LINE_PATTERN.search(line):
        return True
    if re.search(r"\d{1,5}\s+\w+", line) and any(
        part in lowered for part in ["street", "st", "road", "rd", "avenue", "ave", "drive", "dr", "lane", "ln"]
    ):
        return True
    return False


def _normalize_section_title(line: str) -> str | None:
    candidate = re.sub(r"[^A-Za-z\s]", "", line).strip().upper()
    return SECTION_ALIASES.get(candidate)


def _guess_section(line: str) -> str:
    lowered = line.lower()
    if any(token in lowered for token in ["bachelor", "master", "university", "college", "gpa", "degree"]):
        return "EDUCATION"
    if any(token in lowered for token in ["skills", "technical", "microsoft", "sql", "salesforce", "quickbooks", "peoplesoft"]):
        return "TECHNICAL SKILLS"
    if DATE_PATTERN.search(line) or MONTH_DATE_RANGE_PATTERN.search(line):
        return "PROFESSIONAL EXPERIENCE"
    if any(token in lowered for token in ["consultant", "manager", "director", "analyst", "officer"]):
        return "PROFESSIONAL EXPERIENCE"
    return "SUMMARY OF QUALIFICATIONS"


def _set_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _apply_document_theme(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Cambria"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    if "List Bullet" in doc.styles:
        bullet = doc.styles["List Bullet"]
        bullet.font.name = "Cambria"
        bullet.font.size = Pt(10.5)
        bullet.paragraph_format.left_indent = Inches(0.2)
        bullet.paragraph_format.first_line_indent = Inches(-0.12)
        bullet.paragraph_format.space_before = Pt(0)
        bullet.paragraph_format.space_after = Pt(2)
        bullet.paragraph_format.line_spacing = 1.0


def _add_section_header(doc: Document, title: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Cambria"
    run.font.size = Pt(11)
    _set_bottom_border(paragraph)


def _add_left_right_line(doc: Document, left_text: str, right_text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.5))
    left = paragraph.add_run(left_text)
    left.bold = True
    left.font.size = Pt(10.5)
    right = paragraph.add_run(f"\t{right_text}")
    right.italic = True
    right.font.size = Pt(10)


def _document_from_letterhead(letterhead_path: str | None) -> Document:
    """Create a Document based on a letterhead template.

    If a valid `letterhead_path` is provided, the template is loaded and returned
    unchanged so that any header/footer content, images, or styled paragraphs are
    preserved. If the path is missing or invalid, a fresh blank Document is
    created.
    """
    if letterhead_path and Path(letterhead_path).exists():
        doc = Document(letterhead_path)
    else:
        doc = Document()
    # Previously the implementation removed all body elements except the
    # section properties, which unintentionally stripped the visual letterhead
    # content that resides in the body. We now retain the full template content.
    return doc


def _extract_sections(lines: list[str]) -> tuple[str, dict[str, list[str]]]:
    name = ""
    start_index = 0
    for idx, line in enumerate(lines[:8]):
        if _is_probable_name(line):
            name = line.upper()
            start_index = idx + 1
            break

    sections = {section: [] for section in SECTION_ORDER}
    current_section = "SUMMARY OF QUALIFICATIONS"
    explicit_section_found = False
    for line in lines[start_index:]:
        normalized = _normalize_section_title(line)
        if normalized:
            current_section = normalized
            explicit_section_found = True
            continue
        if explicit_section_found:
            sections[current_section].append(line)
            continue

        guessed = _guess_section(line)
        sections[guessed].append(line)

    return name, sections


def _rebalance_sections(sections: dict[str, list[str]]) -> dict[str, list[str]]:
    balanced = {key: list(value) for key, value in sections.items()}
    summary_lines: list[str] = []
    summary_limit = 12

    for idx, line in enumerate(balanced["SUMMARY OF QUALIFICATIONS"]):
        guessed = _guess_section(line)
        if guessed != "SUMMARY OF QUALIFICATIONS":
            balanced[guessed].append(line)
            continue

        # Company/location/date-like lines belong to experience.
        if "," in line and any(state in line for state in [", MD", ", NY", ", DC", ", VA", ", FL"]):
            balanced["PROFESSIONAL EXPERIENCE"].append(line)
            continue

        if MONTH_DATE_RANGE_PATTERN.search(line) or DATE_PATTERN.search(line):
            balanced["PROFESSIONAL EXPERIENCE"].append(line)
            continue

        summary_lines.append(line)

    # Keep summary concise and move trailing non-summary lines if any.
    balanced["SUMMARY OF QUALIFICATIONS"] = summary_lines
    return balanced


def _fallback_name_from_filename(candidate_filename: str) -> str:
    base = candidate_filename.rsplit(".", 1)[0]
    cleaned = re.sub(r"[_\-]+", " ", base)
    cleaned = re.sub(r"\b(resume|cv|std\d+|\d{4})\b", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned.upper() if cleaned else base.upper()



def build_formatted_resume_docx(
    structured_resume: dict,
    letterhead_path: str | None = None,
) -> bytes:
    # 1. Create document from letterhead
    doc = _document_from_letterhead(letterhead_path)
    _apply_document_theme(doc)

    # 2. Candidate name (from JSON, no guessing)
    title = doc.add_paragraph()
    name_run = title.add_run(structured_resume.get("name", ""))
    name_run.bold = True
    name_run.font.name = "Cambria"
    name_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(10)

    # 3. Optional title
    if structured_resume.get("title"):
        subtitle = doc.add_paragraph(structured_resume["title"])
        subtitle.paragraph_format.space_after = Pt(8)

    # 4. Summary
    _add_section_header(doc, "SUMMARY OF QUALIFICATIONS")
    doc.add_paragraph(structured_resume.get("summary", ""))

    # 5. Professional Experience
    _add_section_header(doc, "PROFESSIONAL EXPERIENCE")

    for job in structured_resume.get("professional_experience", []):
        _add_left_right_line(
            doc,
            f"{job['company']} – {job.get('location', '')}",
            job.get("dates", "")
        )

        role_para = doc.add_paragraph(job.get("role", ""))
        role_para.paragraph_format.space_after = Pt(2)

        for bullet in job.get("bullets", []):
            b = doc.add_paragraph(bullet, style="List Bullet")
            b.paragraph_format.space_after = Pt(2)

    # 6. Education
    if structured_resume.get("education"):
        _add_section_header(doc, "EDUCATION")
        for edu in structured_resume["education"]:
            doc.add_paragraph(
                f"{edu.get('institution', '')} – {edu.get('degree', '')}"
            )

    # 7. Technical Skills
    if structured_resume.get("technical_skills"):
        _add_section_header(doc, "TECHNICAL SKILLS")
        doc.add_paragraph(", ".join(structured_resume["technical_skills"]))

    # 8. Certifications
    if structured_resume.get("certifications"):
        _add_section_header(doc, "CERTIFICATIONS")
        for cert in structured_resume["certifications"]:
            doc.add_paragraph(cert)

    # 9. Save
    output = BytesIO()
    doc.save(output)
    return output.getvalue()
