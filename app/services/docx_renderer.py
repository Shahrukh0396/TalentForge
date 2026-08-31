from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Length
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import  WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import random

# =========================================================
# LETTERHEAD PATH
# =========================================================

BASE_DIR = Path(__file__).resolve().parents[2]
LETTERHEAD_PATH = BASE_DIR / "app" / "templates" / "syntax_talent_letterhead.docx"

# =========================================================
# CONFIG
# =========================================================

FONT_NAME = "Cambria"
FONT_SIZE = Pt(10.5)
NAME_SIZE = Pt(14)
COMPANY_SIZE = Pt(11.5)
SECTION_HEADER_SIZE = Pt(11)

PAGE_MARGIN_TOP = Inches(0.5)
PAGE_MARGIN_BOTTOM = Inches(0.5)
PAGE_MARGIN_LEFT = Inches(0.7)
PAGE_MARGIN_RIGHT = Inches(0.7)

SECTION_SPACE_BEFORE = Pt(10)
SECTION_SPACE_AFTER = Pt(4)
FIRST_SECTION_SPACE_BEFORE = Pt(4)
NAME_SPACE_AFTER = Pt(8)
BODY_LINE_SPACING = 1.0
BULLET_LEFT_INDENT = Inches(0.2)
BULLET_HANGING_INDENT = Inches(-0.12)
JOB_BLOCK_SPACE_BEFORE = Pt(6)

# =========================================================
# NORMALIZATION
# =========================================================

NORMALIZATION_MAP = {
    "salesforce": "Salesforce",
    "quickbooks": "QuickBooks",
    "powerpoint": "PowerPoint",
    "peoplesoft": "PeopleSoft",
    "aws": "AWS",
    "gcp": "GCP",
    "azure": "Azure",
}

def ensure_list(v):
    if isinstance(v, list):
        return v
    if v:
        return [v]
    return []


def ensure_bullet_numbering(doc):
    numbering = doc.part.numbering_part.numbering_definitions
    if numbering._numbering is None or len(numbering._numbering) == 0:
        numbering._numbering = OxmlElement("w:numbering")


def right_tab_position(doc: Document):
    section = doc.sections[-1]
    return section.page_width - section.left_margin - section.right_margin


def normalize_paragraph_flow(p):
    pf = p.paragraph_format
    pf.widow_control = False
    pf.keep_with_next = False
    pf.keep_together = False
    if pf.page_break_before:
        pf.page_break_before = False


def apply_page_layout(doc):
    for section in doc.sections:
        section.top_margin = PAGE_MARGIN_TOP
        section.bottom_margin = PAGE_MARGIN_BOTTOM
        section.left_margin = PAGE_MARGIN_LEFT
        section.right_margin = PAGE_MARGIN_RIGHT


def prepare_letterhead_document() -> Document:
    doc = Document(str(LETTERHEAD_PATH))

    # Remove placeholder body paragraphs so we don't carry blank lines/page gaps.
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    return doc


def apply_global_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = BODY_LINE_SPACING


def coerce_length(value) -> Length:
    if isinstance(value, Length):
        return value
    return Pt(value)


def apply_spacing(
    p,
    after: int | Length = 2,
    before: int | Length = 0,
    line_spacing: float = BODY_LINE_SPACING,
):
    pf = p.paragraph_format
    pf.space_before = coerce_length(before)
    pf.space_after = coerce_length(after)
    pf.line_spacing = line_spacing


def add_section_header(doc, title: str, *, first: bool = False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = FIRST_SECTION_SPACE_BEFORE if first else SECTION_SPACE_BEFORE
    pf.space_after = SECTION_SPACE_AFTER
    pf.line_spacing = 1.0

    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = SECTION_HEADER_SIZE

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")

    pBdr.append(bottom)
    pPr.append(pBdr)

    normalize_paragraph_flow(p)
    return p



def add_bullet(doc, text: str):
    text = safe_text(text)
    if not text:
        return

    p = doc.add_paragraph()
    run = p.add_run(f"• {text}")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE

    pf = p.paragraph_format
    pf.left_indent = BULLET_LEFT_INDENT
    pf.first_line_indent = BULLET_HANGING_INDENT
    apply_spacing(p, after=2, before=0)
    normalize_paragraph_flow(p)
    return p


def add_left_right_line(
    doc,
    left_text: str,
    right_text: str = "",
    *,
    left_bold: bool = False,
    left_italic: bool = False,
    left_size=FONT_SIZE,
    right_italic: bool = False,
    space_before: int = 0,
):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(
        right_tab_position(doc), WD_TAB_ALIGNMENT.RIGHT
    )

    left = p.add_run(left_text)
    left.bold = left_bold
    left.italic = left_italic
    left.font.name = FONT_NAME
    left.font.size = left_size

    if right_text:
        p.add_run("\t")
        right = p.add_run(right_text)
        right.font.name = FONT_NAME
        right.font.size = FONT_SIZE
        right.italic = right_italic

    apply_spacing(p, after=2, before=space_before)
    normalize_paragraph_flow(p)
    return p

def safe_text(v):
    return str(v).strip() if v else ""


def _shuffled_competencies(structured_resume: dict) -> list[str]:
    competencies = [
        safe_text(item)
        for item in ensure_list(structured_resume.get("core_competencies"))
        if safe_text(item)
    ]
    random.shuffle(competencies)
    return competencies


def _certification_text(cert) -> str:
    if isinstance(cert, dict):
        return safe_text(
            cert.get("name")
            or cert.get("title")
            or cert.get("certification")
        )
    return safe_text(cert)


def render_resume_docx(structured_resume: dict) -> bytes:
    doc = prepare_letterhead_document()
    ensure_bullet_numbering(doc)
    apply_page_layout(doc)
    apply_global_style(doc)

    # NAME
    name = safe_text(structured_resume.get("name"))
    if not name:
        raise ValueError("Structured resume is missing candidate name")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name.upper())
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = NAME_SIZE
    apply_spacing(p, after=NAME_SPACE_AFTER, before=0)


    # SUMMARY (+ core competencies woven in, no separate section)
    add_section_header(doc, "Summary of Qualifications", first=True)
    summary = structured_resume.get("summary")
    if summary:
        p = doc.add_paragraph()
        summary_run = p.add_run(safe_text(summary))
        summary_run.font.name = FONT_NAME
        summary_run.font.size = FONT_SIZE
        apply_spacing(p, after=4, before=0)

    for competency in _shuffled_competencies(structured_resume):
        add_bullet(doc, competency)

    # PROFESSIONAL EXPERIENCE
    add_section_header(doc, "Professional Experience")
    experience = ensure_list(structured_resume.get("professional_experience"))

    for index, job in enumerate(experience):
        company = safe_text(job.get("company"))
        location = safe_text(job.get("location"))
        role = safe_text(job.get("role"))
        dates = safe_text(job.get("dates"))

        add_left_right_line(
            doc,
            company,
            location,
            left_bold=True,
            left_size=COMPANY_SIZE,
            space_before=JOB_BLOCK_SPACE_BEFORE if index > 0 else 0,
        )

        add_left_right_line(
            doc,
            role,
            dates,
            left_italic=True,
        )

        for b in ensure_list(job.get("bullets")):
            if safe_text(b):
                add_bullet(doc, b)

    # EDUCATION
    education = ensure_list(structured_resume.get("education"))
    if education:
        add_section_header(doc, "Education")
        for edu in education:
            if safe_text(edu):
                p = doc.add_paragraph()
                edu_run = p.add_run(safe_text(edu))
                edu_run.font.name = FONT_NAME
                edu_run.font.size = FONT_SIZE
                apply_spacing(p, after=2, before=0)

    # TECHNICAL SKILLS
    skills = ensure_list(structured_resume.get("technical_skills"))
    clean_skills = [safe_text(s) for s in skills if safe_text(s)]

    # CERTIFICATIONS
    certifications = [
        text
        for cert in ensure_list(structured_resume.get("certifications"))
        if (text := _certification_text(cert))
    ]
    if certifications:
        add_section_header(doc, "Certifications, Licenses & Professional Affiliations")
        for cert in certifications:
            add_bullet(doc, cert)

    if clean_skills:
        add_section_header(doc, "Technical Skills")
        p = doc.add_paragraph()
        skills_run = p.add_run(", ".join(clean_skills))
        skills_run.font.name = FONT_NAME
        skills_run.font.size = FONT_SIZE
        apply_spacing(p, after=4, before=0)

    # SAVE
    out = BytesIO()
    doc.save(out)
    return out.getvalue()